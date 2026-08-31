# -*- coding: utf-8 -*-
"""把批改過的 Word 對照表讀回來，跟原本產出的內容逐格比對。

`make_card_docx.py` 產出的表交給人批改之後，改動可能出現在任何一欄——
不是只有「要改什麼」那欄。所以這裡拿 `tools/out/cards.json`（產表時的原始資料）
當基準，整張表逐格比，把**每一處不一樣的地方**都撈出來。

跑法：python tools/read_card_docx.py
"""
import json
import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "tools/out/球球牌組對照表.docx"
CARDS = ROOT / "tools/out/cards.json"

COLS = ["圖", "牌名", "費用類型", "好認嗎", "說明", "磨爪後", "要改什麼"]


def norm(s: str) -> str:
    """比對前把空白統一：Word 會把換行、全形空白、行首縮排帶進來。"""
    return re.sub(r"\s+", " ", (s or "").replace("　", " ")).strip()


def main() -> None:
    rows = {r["id"]: r for r in json.loads(CARDS.read_text(encoding="utf-8"))}
    doc = Document(DOCX)
    table = doc.tables[0]

    changed_text, changed_up, comments, unknown = [], [], [], []

    for row in list(table.rows)[1:]:
        c = [cell.text for cell in row.cells]
        # 牌名那格是「名稱 / 編號 /（起手牌組 ×N）」三行，編號在第二行
        lines = [x for x in c[1].split("\n") if x.strip()]
        cid = lines[1].strip() if len(lines) > 1 else ""
        src = rows.get(cid)
        if not src:
            unknown.append(norm(c[1])[:40])
            continue
        name = lines[0].strip()

        if norm(c[4]) != norm(src["text"]):
            changed_text.append((cid, name, src["text"], norm(c[4])))
        if norm(c[5]) != norm(src["upgraded"]):
            changed_up.append((cid, name, src["upgraded"], norm(c[5])))
        if norm(c[6]):
            comments.append((cid, name, norm(c[6])))

    print(f"表格共 {len(table.rows) - 1} 列")
    if unknown:
        print(f"認不出編號的列 {len(unknown)}：{unknown}")

    print(f"\n===== 說明文字被改了 {len(changed_text)} 處 =====")
    for cid, name, before, after in changed_text:
        print(f"[{name}] {cid}\n   原：{before}\n   改：{after}")

    print(f"\n===== 磨爪後文字被改了 {len(changed_up)} 處 =====")
    for cid, name, before, after in changed_up:
        print(f"[{name}] {cid}\n   原：{before}\n   改：{after}")

    print(f"\n===== 「要改什麼」欄寫了 {len(comments)} 則 =====")
    for cid, name, text in comments:
        print(f"[{name}] {cid}\n   {text}")

    # 表格後面那幾頁的整批意見
    after_table = []
    hit = False
    for p in doc.paragraphs:
        t = norm(p.text)
        if not t:
            continue
        if t.startswith("怎麼用這份表"):
            hit = True
            continue
        if hit:
            after_table.append(t)
    print(f"\n===== 表格後面的文字 {len(after_table)} 行 =====")
    for t in after_table:
        print(" ", t)


if __name__ == "__main__":
    main()
