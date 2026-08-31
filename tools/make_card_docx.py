# -*- coding: utf-8 -*-
"""把 78 張牌做成一份 Word 對照表，給人逐張批改。

規則文字**不是這裡重寫的**，是 `tools/dump_cards.test.ts` 掛在 vitest 底下、
用遊戲本人那份 `describeCard()` 倒出來的，所以跟畫面上看到的一字不差。

跑法：
  npx vitest run tools/dump_cards.test.ts      # 先產 tools/out/cards.json
  python tools/make_card_docx.py               # 再產 docx

輸出：`tools/out/球球牌組對照表.docx`

表格欄位刻意留了最右邊一欄空白給人寫意見——這份文件的用途就是拿來改的。
"""
import io
import json
from pathlib import Path

import numpy as np
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CARDS = ROOT / "tools/out/cards.json"
MANIFEST = ROOT / "public/assets/manifest.json"
OUT = ROOT / "tools/out/球球牌組對照表.docx"

# 牌面在 Word 裡的寬度。3.2 公分約等於遊戲裡看到的大小，
# 太小看不出圖畫了什麼，太大整份會變成一百多頁。
ART_W = Cm(3.2)

TYPE_COLOUR = {"攻擊": RGBColor(0xB8, 0x34, 0x2A),
               "技能": RGBColor(0x2E, 0x6D, 0xB4),
               "能力": RGBColor(0x7A, 0x3F, 0xAE)}


def as_png(path: Path) -> io.BytesIO:
    """WebP 轉 PNG。python-docx 認不得 WebP（會丟 UnrecognizedImageError），
    所以每張都在記憶體裡轉一次，不落地成暫存檔。
    順便把透明底合成白色：牌面在遊戲裡就是白底，透明的放進 Word 會跟頁面同色、看不出邊界。"""
    im = Image.open(path).convert("RGBA")
    flat = Image.new("RGB", im.size, (255, 255, 255))
    flat.paste(im, mask=im.split()[3])
    buf = io.BytesIO()
    flat.save(buf, "PNG")
    buf.seek(0)
    return buf


def distinct_scores(manifest: dict) -> dict[str, float]:
    """每張牌跟「最像的三張」的平均差距。越小＝越容易跟別張搞混。

    跟 `tools/card_distinct.py` 用同一套算法：縮到 16x16 再比，比的是一眼掃過去的印象。
    這裡把分數印進表格，是為了讓看的人知道**哪幾張該優先重畫**——
    不然七十八張看下來只會覺得「好像都還好」。
    """
    ids, sigs = [], []
    for key, rel in sorted(manifest["cards"].items()):
        im = Image.open(ROOT / "public" / rel).convert("RGBA").resize((16, 16), Image.LANCZOS)
        a = np.asarray(im).astype(float)
        al = a[:, :, 3] / 255
        ids.append(key.replace("card/", ""))
        sigs.append(np.concatenate([(a[:, :, i] * al).ravel() for i in range(3)]))
    m = np.array(sigs)
    d = np.abs(m[:, None, :] - m[None, :, :]).mean(axis=2)
    np.fill_diagonal(d, 1e9)
    near = np.sort(d, axis=1)[:, :3].mean(axis=1)
    return dict(zip(ids, near))


def set_cell_bg(cell, hex_colour: str) -> None:
    """python-docx 沒有現成的儲存格底色 API，只能自己塞 XML。"""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def main() -> None:
    rows = json.loads(CARDS.read_text(encoding="utf-8"))
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    scores = distinct_scores(manifest)
    # 分數最低的四分之一標紅：那些就是「畫面沒特色、容易跟別張混在一起」的
    weak_at = float(np.percentile(list(scores.values()), 25))

    doc = Document()
    # 橫式：欄位多，直式會把說明擠成一行三個字
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Cm(1.4))

    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(10)
    # 中文字要另外指定東亞字型，不然 Word 會用預設的細明體
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    h = doc.add_heading("球球勇闖魔物塔　牌組對照表", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(f"共 {len(rows)} 張。").bold = True
    intro.add_run(
        "規則文字是從遊戲程式直接倒出來的，跟畫面上看到的一字不差。"
        "最右邊一欄留白，請直接在上面寫要改什麼——圖不對、字不順、效果想調都可以。")

    note = doc.add_paragraph()
    note.add_run("「好認嗎」那一欄").bold = True
    note.add_run(
        "是量出來的：把每張牌縮成一眼掃過去的印象，算它跟最像的三張差多少。"
        f"數字越小越容易跟別張搞混，低於 {weak_at:.0f} 的會標成紅色的「該重畫」。"
        "這副牌全部 78 張的平均是 "
        f"{np.mean(list(scores.values())):.1f}，最差 {min(scores.values()):.1f}。")

    doc.add_paragraph()

    # 依「牌庫 → 類型 → 費用」分組，同一組的牌會排在一起，比較好一次看完
    order = {"起手": 0, "忍術": 1, "絕學": 2, "壞毛病": 3}
    rows.sort(key=lambda r: (order.get(r["pool"], 9), r["type"], r["cost"], r["name"]))

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(3.5), Cm(3.8), Cm(2.6), Cm(2.2), Cm(5.6), Cm(4.8), Cm(5.6)]
    heads = ["牌面圖", "牌名 / 編號", "費用 / 類型", "好認嗎", "現在的說明", "磨爪之後", "要改什麼（請寫這欄）"]

    hdr = table.rows[0]
    for i, (cell, text) in enumerate(zip(hdr.cells, heads)):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(cell, "3A2A1C")
        run.font.color.rgb = RGBColor(0xFF, 0xEE, 0xCA)
        cell.width = widths[i]

    for r in rows:
        row = table.add_row()
        c = row.cells
        for i in range(7):
            c[i].width = widths[i]

        # 圖
        rel = manifest["cards"].get(r["art"])
        p = c[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if rel and (ROOT / "public" / rel).exists():
            p.add_run().add_picture(as_png(ROOT / "public" / rel), width=ART_W)
        else:
            p.add_run("（缺圖）")

        # 牌名與編號
        p = c[1].paragraphs[0]
        run = p.add_run(r["name"])
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = TYPE_COLOUR.get(r["type"], RGBColor(0x2B, 0x21, 0x18))
        sub = c[1].add_paragraph()
        idrun = sub.add_run(r["id"])
        idrun.font.size = Pt(8)
        idrun.font.color.rgb = RGBColor(0x88, 0x80, 0x76)
        if r["starter"]:
            s = c[1].add_paragraph()
            srun = s.add_run(f"起手牌組 ×{r['starter']}")
            srun.font.size = Pt(8.5)
            srun.bold = True

        # 費用 / 類型 / 稀有 / 牌庫
        c[2].text = ""
        for line, size in ((f"{r['cost']} 飯糰", 10), (r["type"], 10),
                           (r["rarity"], 9), (r["pool"], 9)):
            pp = c[2].paragraphs[0] if c[2].paragraphs[0].text == "" and size == 10 and line.endswith("飯糰") \
                else c[2].add_paragraph()
            rr = pp.add_run(line)
            rr.font.size = Pt(size)
            if size == 9:
                rr.font.color.rgb = RGBColor(0x88, 0x80, 0x76)

        # 好認嗎
        sc = scores.get(r["id"])
        c[3].text = ""
        pp = c[3].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sc is None:
            pp.add_run("—")
        else:
            rr = pp.add_run(f"{sc:.0f}")
            rr.bold = True
            rr.font.size = Pt(13)
            if sc < weak_at:
                rr.font.color.rgb = RGBColor(0xC0, 0x28, 0x20)
                warn = c[3].add_paragraph()
                warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
                wr = warn.add_run("該重畫")
                wr.font.size = Pt(8)
                wr.font.color.rgb = RGBColor(0xC0, 0x28, 0x20)
                set_cell_bg(c[3], "FBE5E3")
            else:
                rr.font.color.rgb = RGBColor(0x3A, 0x6B, 0x35)

        c[4].text = r["text"]
        c[5].text = r["upgraded"]
        c[6].text = ""

    doc.add_page_break()
    doc.add_heading("怎麼用這份表", level=1)
    for line in [
        "最右邊那欄直接寫，想到什麼寫什麼，不用管格式。",
        "「圖不對」請描述你想看到的畫面，例如「這張應該要有兩隻貓」「顏色太暗」。",
        "「字不順」直接把你想要的句子寫上去，我照抄。",
        "「效果想調」寫清楚改成多少，例如「6 點傷害改成 5 點」。",
        "整批想調的（例如「所有攻擊牌的圖都太紅」）寫在最後面那頁就好，不用每張都寫。",
        "標紅色「該重畫」的那幾張，多半是圖裡只有貓在擺姿勢、沒有能一眼認出來的東西。"
        "如果你有想法（要放什麼道具、什麼顏色）寫上去最好；沒想法也可以只寫「重畫」。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"{len(rows)} 張 → {OUT}")
    print(f"檔案大小 {OUT.stat().st_size / 1024 / 1024:.1f} MB")


if __name__ == "__main__":
    main()
